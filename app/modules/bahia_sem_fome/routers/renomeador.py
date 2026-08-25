import re
import io
import json
import logging
import zipfile
import asyncio
from typing import List
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
import fitz # PyMuPDF
import pdfplumber
import os
from docx import Document
from app.config import settings

router = APIRouter(prefix="/api/bahia-sem-fome", tags=["BSF API"])
logger = logging.getLogger(__name__)

class RenameInfo(BaseModel):
    nome: str = Field(description="Nome completo do beneficiário principal (titular da família)")
    tipo: str = Field(description="Tipo de documento: ATESTE ou COLLETUM")
    cpf: str = Field(description="CPF do beneficiário, formatado com pontos e hífen (ex: '000.000.000-00'), ou vazio se não encontrado")
    atividade: str = Field(description="Descrição resumida da atividade que está assinalada/marcada com um 'X' (ou circulada, assinalada de qualquer forma) na tabela/lista de TIPO DE ATIVIDADE, em maiúsculas e sem acentos")
    data: str = Field(description="Data da atividade escrita no documento no formato DD-MM-AAAA ou vazio se não encontrada")
    tecnico: str = Field(description="Nome do técnico(a) responsável indicado no campo 'TÉCNICO(A)' ou 'TÉCNICO', em maiúsculas e sem acentos, ou vazio se não encontrado")
    comunidade: str = Field(description="Nome da localidade/comunidade indicada no campo 'COMUNIDADE' ou 'LOCAL', em maiúsculas e sem acentos, ou vazio se não encontrada")

MODELOS_PERMITIDOS = [
    "gemini-2.5-flash-lite",
    "gemini-2.5-flash",
    "gemini-3.5-flash-lite",
    "gemini-3.1-flash-lite",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-1.5-flash-8b",
    "gemini-3.5-flash",
    "gemini-3.6-flash"
]

_current_model_index = 0

def obter_modelos_ordenados() -> List[str]:
    """Retorna a lista de modelos ordenados a partir do modelo atual no rodízio (Round-Robin/Load Balancer)."""
    global _current_model_index
    _current_model_index = (_current_model_index + 1) % len(MODELOS_PERMITIDOS)
    return MODELOS_PERMITIDOS[_current_model_index:] + MODELOS_PERMITIDOS[:_current_model_index]

def get_gemini_client():
    api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
    if api_key:
        return genai.Client(api_key=api_key)
    return None

def pdf_page_to_png_bytes(pdf_content: bytes, page_num: int = 0) -> bytes:
    """Renderiza uma página do PDF para imagem PNG em memória usando PyMuPDF (fitz)."""
    with fitz.open(stream=pdf_content, filetype="pdf") as doc:
        if page_num >= len(doc):
            page_num = 0
        page = doc[page_num]
        # Renderiza a página com DPI 150 para equilíbrio entre legibilidade da IA e consumo de dados
        pix = page.get_pixmap(dpi=150)
        return pix.tobytes("png")

def extrair_local_regex(texto: str):
    nome, tipo = None, None
    texto_flat = re.sub(r'\s+', ' ', texto.upper())
    letras_br = r"[A-ZÁÉÍÓÚÀÈÌÒÙÂÊÎÔÛÃÕÇ\s\'\-]"
    
    # A Tesoura: Corta o texto assim que bater em um número ou rótulo do PDF
    paradas = r'(\d|CPF|DAP|CAF|RG|DATA|MUNIC|ENDERE|P[AÁ]GINA)'

    # PADRÃO 1: Ateste de Atividade
    if "ATESTE" in texto_flat:
        tipo = "ATESTE"
        # Pega um bloco sujo de até 150 caracteres após a âncora
        match = re.search(r'BENEFICI[AÁ]RI[OA]?\s*\(?A?\)?\s*[:\-]?\s*(.{5,150})', texto_flat)
        if match:
            bloco = match.group(1)
            fatia = re.split(paradas, bloco)[0] # Corta na primeira parada
            nome_limpo = re.sub(r'[^A-ZÁÉÍÓÚÀÈÌÒÙÂÊÎÔÛÃÕÇ\s\'\-]', '', fatia).strip()
            nome = re.sub(r'\s+', ' ', nome_limpo)

    # PADRÃO 2: Formulário de Atividade (Colletum)
    elif "FORMUL" in texto_flat or "COLLETUM" in texto_flat:
        tipo = "COLLETUM"
        # Tenta pegar direto entre CPFs primeiro
        match = re.search(r'\d{3}\.?\d{3}\.?\d{3}-?\d{2}\s*(' + letras_br + r'{5,120})\s*(?:CPF|DATA|\d{3})', texto_flat)
        if match:
            nome = match.group(1).strip()
        # Fallback Colletum 
        if not nome:
            match = re.search(r'NOME DO BENEFICI[AÁ]RI[OA]?\s*\(?A?\)?\s*[:\-]?\s*(.{5,150})', texto_flat)
            if match:
                bloco = match.group(1)
                fatia = re.split(paradas, bloco)[0]
                nome_limpo = re.sub(r'[^A-ZÁÉÍÓÚÀÈÌÒÙÂÊÎÔÛÃÕÇ\s\'\-]', '', fatia).strip()
                nome = re.sub(r'\s+', ' ', nome_limpo)

    # PADRÃO 3: Acompanhamento Plano Produtivo
    elif "ACOMPANHAMENTO PLANO PRODUTIVO" in texto_flat:
        tipo = "ACOMPANHAMENTO"
        match = re.search(r'TITULAR 1 DO GRUPO FAMILIAR\s*[:\-]?\s*(.{5,150})', texto_flat)
        if match:
            bloco = match.group(1)
            fatia = re.split(paradas, bloco)[0]
            nome_limpo = re.sub(r'[^A-ZÁÉÍÓÚÀÈÌÒÙÂÊÎÔÛÃÕÇ\s\'\-]', '', fatia).strip()
            nome = re.sub(r'\s+', ' ', nome_limpo)

        if not nome or len(nome) < 5:
            match = re.search(r'NOME\s*:\s*(?:P[AÁ]GINA\s+\d+\s+DE\s+\d+\s*)?(' + letras_br + r'{5,120})\s*CPF DO TITULAR', texto_flat)
            if match:
                nome = match.group(1).strip()

    if nome and len(nome) < 5:
        nome = None

    return nome, tipo

async def extrair_e_analisar(file_content: bytes, filename: str, mode: str = "ateste"):
    """Extrai visualmente (imagem) ou texto do PDF em cascata (Waterfall) e identifica nome, tipo, cpf, atividade e data."""
    try:
        # 1. TENTATIVA: IA Gemini Multimodal (Visão de Imagem) - Método Principal de Alta Precisão
        client = get_gemini_client()
        if client:
            try:
                # Detecta se é imagem direta (JPEG/PNG/BMP) ou PDF
                if file_content.startswith(b'\xff\xd8'):
                    img_bytes = file_content
                    mime_type = "image/jpeg"
                elif file_content.startswith(b'\x89PNG'):
                    img_bytes = file_content
                    mime_type = "image/png"
                else:
                    try:
                        img_bytes = pdf_page_to_png_bytes(file_content, 0)
                        mime_type = "image/png"
                    except Exception:
                        # Se não for PDF válido, tenta usar como imagem genérica
                        img_bytes = file_content
                        mime_type = "image/jpeg"

                image_part = types.Part.from_bytes(
                    data=img_bytes,
                    mime_type=mime_type
                )
                
                if mode == "ateste":
                    prompt = (
                        "Você é um assistente especialista em processamento de documentos do projeto 'Bahia Sem Fome'.\n"
                        "Analise a imagem da página do documento de Ateste fornecida.\n"
                        "Você deve identificar e extrair as seguintes informações:\n"
                        "1. O nome completo do beneficiário principal (titular da família), localizado no campo 'BENEFICIÁRIO(A)' ou 'NOME'.\n"
                        "2. O tipo de documento: Deve ser 'ATESTE'.\n"
                        "3. O CPF do beneficiário principal, localizado no campo 'CPF' ou 'CPF BENEFICIÁRIO'. Retorne formatado com pontos e hífen (ex: '123.456.789-00'). Se não encontrar, retorne vazio.\n"
                        "4. Qual atividade específica está assinalada ou marcada com um 'X' (ou circulada, marcada de qualquer outra forma) na tabela/lista de 'TIPO DE ATIVIDADE'.\n"
                        "   Retorne uma descrição curta e resumida da atividade em maiúsculas e sem acentos (ex: 'VISITA TECNICA', 'CADASTRO', 'CARACTERIZACAO UPF I', 'LEVANTAMENTO SOCIOECONOMICO', etc.).\n"
                        "5. A data da atividade, geralmente preenchida à mão na linha da tabela do cabeçalho sob a coluna 'DATA'.\n"
                        "   Formate a data obrigatoriamente como DD-MM-AAAA (ex: '15-08-2026'). Se não encontrar ou não estiver preenchida, retorne vazio.\n"
                        "6. O nome do Técnico(a) responsável, localizado no campo 'TÉCNICO(A)' ou 'TÉCNICO' no cabeçalho. Retorne em maiúsculas e sem acentos.\n"
                        "7. O nome da Comunidade / Localidade, localizado no campo 'COMUNIDADE' ou 'LOCAL' no cabeçalho. Retorne em maiúsculas e sem acentos.\n\n"
                        "Retorne APENAS um JSON no formato estrito: {\"nome\": \"NOME\", \"tipo\": \"TIPO\", \"cpf\": \"CPF\", \"atividade\": \"ATIVIDADE\", \"data\": \"DATA\", \"tecnico\": \"TECNICO\", \"comunidade\": \"COMUNIDADE\"}."
                    )
                else: # colletum ou auto
                    prompt = (
                        "Você é um assistente especialista em processamento de documentos do projeto 'Bahia Sem Fome'.\n"
                        "Analise a imagem da página do documento fornecido (pode ser um COLLETUM ou ATESTE).\n"
                        "Você deve identificar e extrair as seguintes informações:\n"
                        "1. O nome completo do beneficiário principal (titular da família), localizado no campo de nome/beneficiário.\n"
                        "2. O tipo de documento: Se for um formulário Colletum, retorne 'COLLETUM'. Se for um Ateste, retorne 'ATESTE'.\n"
                        "3. O CPF do beneficiário principal. Retorne formatado com pontos e hífen (ex: '123.456.789-00'). Se não encontrar, retorne vazio.\n"
                        "4. A atividade descrita no formulário. Retorne uma descrição curta em maiúsculas (ex: 'FORMULARIO COLLETUM', 'CADASTRO', etc.).\n"
                        "5. A data do documento. Formate obrigatoriamente como DD-MM-AAAA ou vazio se não encontrar.\n"
                        "6. O nome do Técnico(a) responsável, se indicado no formulário. Retorne em maiúsculas e sem acentos.\n"
                        "7. O nome da Comunidade / Localidade, se indicado no formulário. Retorne em maiúsculas e sem acentos.\n\n"
                        "Retorne APENAS um JSON no formato estrito: {\"nome\": \"NOME\", \"tipo\": \"TIPO\", \"cpf\": \"CPF\", \"atividade\": \"ATIVIDADE\", \"data\": \"DATA\", \"tecnico\": \"TECNICO\", \"comunidade\": \"COMUNIDADE\"}."
                    )

                # Tenta até 2 vezes (com pausa de 60s entre elas se der 429)
                for tentativa_batch in range(2):
                    for nome_modelo in obter_modelos_ordenados():
                        try:
                            logger.info(f"Tentando analisar visualmente '{filename}' com modelo: {nome_modelo} no modo: {mode}")
                            response = await asyncio.to_thread(
                                client.models.generate_content,
                                model=nome_modelo,
                                contents=[image_part, prompt],
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json",
                                    response_schema=RenameInfo
                                )
                            )
                            
                            # Parse do JSON
                            data = json.loads(response.text)
                            nome = data.get("nome", "DESCONHECIDO").strip().upper()
                            tipo = data.get("tipo", "DOCUMENTO").strip().upper()
                            cpf = data.get("cpf", "").strip()
                            atividade = data.get("atividade", "").strip().upper()
                            data_doc = data.get("data", "").strip()
                            tecnico = data.get("tecnico", "").strip().upper()
                            comunidade = data.get("comunidade", "").strip().upper()
                            
                            # Saneamento dos dados
                            nome_saneado = "".join(c for c in nome if c.isalnum() or c in (" ", "-", "_")).strip()
                            cpf_saneado = "".join(c for c in cpf if c.isalnum() or c in (".", "-")).strip()
                            atividade_saneada = "".join(c for c in atividade if c.isalnum() or c in (" ", "-", "_")).strip()
                            data_saneada = "".join(c for c in data_doc if c.isalnum() or c == "-").strip()
                            tecnico_saneado = "".join(c for c in tecnico if c.isalnum() or c in (" ", "-", "_")).strip()
                            comunidade_saneada = "".join(c for c in comunidade if c.isalnum() or c in (" ", "-", "_")).strip()
                            
                            # Monta o novo nome com traços
                            parts = [nome_saneado]
                            if atividade_saneada:
                                parts.append(atividade_saneada)
                            else:
                                parts.append(tipo)
                            
                            if data_saneada:
                                # Garante hífen no lugar de barras ou pontos na data
                                data_saneada = data_saneada.replace("/", "-").replace(".", "-")
                                parts.append(data_saneada)
                                
                            new_name = " - ".join(parts) + ".pdf"
                            logger.info(f"✅ Sucesso visual com modelo '{nome_modelo}' para {filename}: {new_name}")
                            return new_name, {
                                "nome": nome_saneado,
                                "cpf": cpf_saneado,
                                "tipo": tipo,
                                "atividade": atividade_saneada,
                                "data": data_saneada,
                                "tecnico": tecnico_saneado,
                                "comunidade": comunidade_saneada
                            }

                        except Exception as e:
                            err_str = str(e)
                            logger.warning(f"⚠️ Modelo '{nome_modelo}' atingiu limite ou falhou ({err_str[:120]}). Alternando instantaneamente para o próximo modelo do pool...")
                            continue
                            
                    # Se todos os modelos falharem, verifica se é a primeira tentativa.
                    if tentativa_batch == 0:
                        logger.warning(f"⏳ Todos os modelos falharam (provável limite de cota 429 atingido). Pausando a fila por 60 segundos antes de tentar novamente para o arquivo {filename}...")
                        await asyncio.sleep(60)
                    else:
                        logger.error(f"❌ Falha definitiva na IA para {filename} após pausa de recuperação.")
            except Exception as e_img:
                logger.error(f"Erro no processamento visual com Gemini para {filename}: {e_img}")
        else:
            logger.warning("API do Gemini não configurada ou indisponível. Usando fallback local.")

        # 2. TENTATIVA (FALLBACK): Leitura de Texto Local (PyMuPDF / pdfplumber) + Regex
        # Usado apenas se a IA do Gemini falhar ou não estiver configurada.
        nome_local, tipo_local = None, None
        try:
            with fitz.open(stream=file_content, filetype="pdf") as pdf:
                text_fitz = ""
                for i in range(min(3, len(pdf))):
                    text_fitz += pdf[i].get_text() + "\n"
                if text_fitz.strip():
                    nome_local, tipo_local = extrair_local_regex(text_fitz)
        except Exception as e:
            logger.warning(f"Erro no fallback PyMuPDF para {filename}: {e}")

        if not nome_local or not tipo_local:
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    text_plumber = ""
                    for i in range(min(3, len(pdf.pages))):
                        text_plumber += pdf.pages[i].extract_text() + "\n"
                    if text_plumber.strip():
                        nome_local, tipo_local = extrair_local_regex(text_plumber)
            except Exception as e:
                logger.warning(f"Erro no fallback pdfplumber para {filename}: {e}")

        if nome_local and tipo_local:
            nome_saneado = "".join(c for c in nome_local if c.isalnum() or c in (" ", "-", "_")).strip()
            new_name = f"{nome_saneado} - {tipo_local}.pdf"
            logger.info(f"⚡ Sucesso no Fallback Local (Regex) para {filename}: {new_name}")
            return new_name, {
                "nome": nome_saneado,
                "cpf": "",
                "tipo": tipo_local,
                "atividade": "",
                "data": "",
                "tecnico": "",
                "comunidade": ""
            }

        logger.error(f"❌ Todos os métodos falharam para o arquivo {filename}. Mantendo original.")
        return filename, None

    except Exception as e:
        logger.error(f"Erro fatal ao processar {filename}: {e}")
        return filename, None

@router.post("/renomeador-individual")
async def renomear_individual(file: UploadFile = File(...), mode: str = "ateste"):
    """Recebe um único PDF, renomeia-o via IA (ou fallback) e retorna o novo nome."""
    try:
        content = await file.read()
        new_filename, data = await extrair_e_analisar(content, file.filename, mode)
        
        if not new_filename or new_filename.strip() == "":
            new_filename = file.filename
            
        return {"new_name": new_filename, "data": data}
    except Exception as e:
        logger.error(f"Erro ao processar arquivo individual {file.filename}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/renomeador-lote")
async def renomear_lote(files: List[UploadFile] = File(...), mode: str = "ateste"):
    """Recebe múltiplos PDFs, renomeia-os via IA e retorna um ZIP organizado em estrutura de pastas por Técnico/Comunidade/Beneficiário."""
    if not files:
        raise HTTPException(status_code=400, detail="Nenhum arquivo enviado.")

    zip_buffer = io.BytesIO()
    
    try:
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            for file in files:
                content = await file.read()
                
                # Processamento assíncrono para cada arquivo
                new_filename, data = await extrair_e_analisar(content, file.filename, mode)
                
                # Se o nome falhar ou for vazio, usa o original
                if not new_filename or new_filename.strip() == "":
                    new_filename = file.filename
                
                # Estrutura hierárquica de pastas dentro do ZIP: TECNICO / COMUNIDADE / BENEFICIARIO / ARQUIVO.pdf
                if data:
                    tecnico = data.get("tecnico") or "SEM_TECNICO"
                    comunidade = data.get("comunidade") or "SEM_COMUNIDADE"
                    nome = data.get("nome") or "DESCONHECIDO"
                    zip_path = f"{tecnico}/{comunidade}/{nome}/{new_filename}"
                else:
                    zip_path = new_filename
                
                # Adiciona ao ZIP na estrutura de pastas
                zip_file.writestr(zip_path, content)

        zip_buffer.seek(0)
        
        return StreamingResponse(
            iter([zip_buffer.getvalue()]),
            media_type="application/x-zip-compressed",
            headers={
                "Content-Disposition": f"attachment; filename=BSF_Renomeados_{len(files)}_arquivos.zip"
            }
        )

    except Exception as e:
        logger.error(f"Erro na geração do ZIP: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao gerar pacote ZIP: {str(e)}")

class AtesteItem(BaseModel):
    nome: str
    cpf: str
    atividade: str
    data: str

class FichaRequest(BaseModel):
    items: List[AtesteItem]

@router.post("/gerar-ficha-recebimento")
async def gerar_ficha_recebimento(request: FichaRequest):
    """Lê o template de Recebimento de documento.docx, preenche a tabela com os atestes processados e retorna o arquivo."""
    try:
        template_path = settings.BASE_DIR / "app" / "modules" / "bahia_sem_fome" / "assets" / "Recebimento de documento.docx"
        if not template_path.exists():
            raise HTTPException(status_code=404, detail="Template de recebimento não encontrado.")
            
        doc = Document(str(template_path))
        
        # O arquivo tem duas tabelas (Table 0 e Table 1) que contêm colunas NOME, CPF, ATIVIDADE, DATA
        if len(doc.tables) >= 1:
            table0 = doc.tables[0]
            max_t0_rows = len(table0.rows) - 1 # desconsidera cabeçalho
            
            table1 = doc.tables[1] if len(doc.tables) >= 2 else None
            max_t1_rows = len(table1.rows) - 1 if table1 else 0
            
            for idx, item in enumerate(request.items):
                if idx < max_t0_rows:
                    row = table0.rows[idx + 1]
                    row.cells[0].text = item.nome
                    row.cells[1].text = item.cpf
                    row.cells[2].text = item.atividade
                    row.cells[3].text = item.data
                else:
                    t1_idx = idx - max_t0_rows
                    if table1 and t1_idx < max_t1_rows:
                        row = table1.rows[t1_idx + 1]
                        row.cells[0].text = item.nome
                        row.cells[1].text = item.cpf
                        row.cells[2].text = item.atividade
                        row.cells[3].text = item.data
                        
        # Salva o documento gerado em um buffer em memória
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=Recebimento_de_documento.docx"
            }
        )
    except Exception as e:
        logger.error(f"Erro ao gerar ficha de recebimento: {e}")
        raise HTTPException(status_code=500, detail=f"Erro interno: {str(e)}")
