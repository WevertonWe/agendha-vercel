import os
import re
import sys
import shutil
import unicodedata
from pathlib import Path
from typing import List, Set, Optional, Dict, Any

# Garante suporte a caracteres especiais no terminal Windows
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')

try:
    import fitz  # PyMuPDF
except ImportError:
    print("ERRO: PyMuPDF (fitz) não está instalado no ambiente.")
    sys.exit(1)

# Caminho Base dos Técnicos
BASE_TECNICOS_DIR = Path(r"C:\Users\CLIENTE\Desktop\BAHIA_SEM_FOME\weverton\técnicos")

# Mapeamento de nomes de técnicos para suas pastas
MAPA_PASTAS_TECNICOS = {
    "CAROLINE": "caroline",
    "CRISTINA": "cristina",
    "JOSEFA CRISTINA": "cristina",
    "ESTELLA": "estella",
    "MIGSON": "migson",
    "TIAGO": "thiago",
    "THIAGO": "thiago",
    "LUIZ ANTONIEL": "tony",
    "TONY": "tony",
    "WANDISSON": "Wandisson"
}

def normalizar_texto(texto: str) -> str:
    """Remove acentos, pontuações e converte para maiúsculas para comparação segura."""
    if not texto:
        return ""
    nfkd_form = unicodedata.normalize('NFKD', str(texto))
    texto_sem_acento = "".join([c for c in nfkd_form if not unicodedata.combining(c)])
    limpo = re.sub(r'[^A-Z0-9\s]', '', texto_sem_acento.upper())
    return re.sub(r'\s+', ' ', limpo).strip()

def normalizar_cpf(cpf: str) -> str:
    """Extrai apenas os números do CPF."""
    if not cpf:
        return ""
    return re.sub(r'\D', '', str(cpf))

def normalizar_data(data_str: str) -> str:
    """Normaliza datas para o formato DD.MM.AAAA."""
    if not data_str:
        return "SEM_DATA"
    parts = data_str.replace('/', '-').replace('.', '-').split('-')
    if len(parts) == 3:
        if len(parts[2]) == 4:
            return f"{int(parts[0]):02d}.{int(parts[1]):02d}.{parts[2]}"
        elif len(parts[0]) == 4:
            return f"{int(parts[2]):02d}.{int(parts[1]):02d}.{parts[0]}"
    return data_str.replace('/', '.')

def extrair_dados_pdf(pdf_path: Path) -> Dict[str, str]:
    """Lê o PDF do Coletum e extrai Beneficiário, CPF, Comunidade, Técnico e Data da Atividade."""
    doc = fitz.open(pdf_path)
    text = ""
    for i in range(min(3, len(doc))):
        text += doc[i].get_text() + "\n"
    doc.close()

    # Beneficiário
    m_nome = re.search(r'Dados do Grupo Familiar\s*\n\s*Nome\s*\n\s*([^\n]+)', text, re.IGNORECASE)
    if not m_nome:
        m_nome = re.search(r'Nome\s*:\s*([^\n]+)', text, re.IGNORECASE)
    if not m_nome:
        m_nome = re.search(r'BENEFICI[AÁ]RI[OA]?\s*[:\-]?\s*([^\n]+)', text, re.IGNORECASE)
    nome = m_nome.group(1).strip() if m_nome else "DESCONHECIDO"

    # CPF
    m_cpf = re.search(r'Dados do Grupo Familiar\s*\n.*?\n\s*CPF\s*\n\s*([\d\.\-]+)', text, re.DOTALL | re.IGNORECASE)
    if not m_cpf:
        m_cpf = re.search(r'(\d{3}\.\d{3}\.\d{3}\-\d{2})', text)
    cpf = m_cpf.group(1).strip() if m_cpf else ""

    # Comunidade
    m_com = re.search(r'Comunidade\s*\n\s*([^\n]+)', text, re.IGNORECASE)
    comunidade = m_com.group(1).strip() if m_com else "GERAL"

    # Técnico
    m_tec = re.search(r'Nome do\(a\) t[eé]cnico\(a\)\s*respons[aá]vel\s*\n\s*([^\n]+)', text, re.IGNORECASE)
    tecnico = m_tec.group(1).strip() if m_tec else ""

    # Data da atividade
    m_data = re.search(r'Data da realiza[çc][ãa]o da\s*\n\s*atividade\s*\n\s*(\d{2}/\d{2}/\d{4})', text, re.IGNORECASE)
    if not m_data:
        m_data = re.search(r'(\d{2}/\d{2}/\d{4})', text)
    data = m_data.group(1).strip() if m_data else ""

    return {
        "nome": nome,
        "cpf": cpf,
        "comunidade": comunidade,
        "tecnico": tecnico,
        "data": data,
        "data_formatada": normalizar_data(data)
    }

def carregar_lista_permitidos(fonte: Any) -> Set[str]:
    """
    Carrega a lista de beneficiários autorizados a serem movidos.
    Aceita: lista de strings, caminho para arquivo .txt, .docx ou texto bruto.
    """
    permitidos = set()

    if isinstance(fonte, (list, set, tuple)):
        for item in fonte:
            norm = normalizar_texto(item)
            if norm: permitidos.add(norm)
            cpf_limpo = normalizar_cpf(item)
            if len(cpf_limpo) == 11: permitidos.add(cpf_limpo)

    elif isinstance(fonte, (str, Path)):
        p = Path(fonte)
        if p.exists() and p.is_file():
            if p.suffix.lower() == '.docx':
                import docx
                doc = docx.Document(p)
                for para in doc.paragraphs:
                    if para.text.strip():
                        permitidos.add(normalizar_texto(para.text))
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            txt = cell.text.strip()
                            if txt:
                                permitidos.add(normalizar_texto(txt))
                                cpf_limpo = normalizar_cpf(txt)
                                if len(cpf_limpo) == 11: permitidos.add(cpf_limpo)
            else:
                # Arquivo de texto
                with open(p, 'r', encoding='utf-8', errors='ignore') as f:
                    for line in f:
                        line_clean = line.strip()
                        if line_clean:
                            permitidos.add(normalizar_texto(line_clean))
                            cpf_limpo = normalizar_cpf(line_clean)
                            if len(cpf_limpo) == 11: permitidos.add(cpf_limpo)
        else:
            # Texto passado como string
            for line in str(fonte).strip().splitlines():
                parts = re.split(r'[\t;,|]', line)
                for part in parts:
                    part_clean = part.strip()
                    if part_clean:
                        permitidos.add(normalizar_texto(part_clean))
                        cpf_limpo = normalizar_cpf(part_clean)
                        if len(cpf_limpo) == 11: permitidos.add(cpf_limpo)

    return permitidos

def encontrar_pasta_tecnico(tecnico_nome: str) -> Optional[Path]:
    """Identifica a pasta raiz do técnico."""
    norm_tec = normalizar_texto(tecnico_nome)
    for chave, pasta_nome in MAPA_PASTAS_TECNICOS.items():
        if chave in norm_tec:
            pasta = BASE_TECNICOS_DIR / pasta_nome
            if pasta.exists():
                return pasta
    # Fallback: procura se o nome de alguma pasta está contido no nome do técnico
    for sub in BASE_TECNICOS_DIR.iterdir():
        if sub.is_dir() and normalizar_texto(sub.name) in norm_tec:
            return sub
    return None

def encontrar_ou_criar_pasta_beneficiario(pasta_tecnico: Path, comunidade: str, beneficiario: str) -> Path:
    """Localiza a pasta da comunidade e do beneficiário na estrutura do técnico."""
    doc_ativ = pasta_tecnico / "documentos-atividades"
    doc_ativ.mkdir(parents=True, exist_ok=True)

    norm_com = normalizar_texto(comunidade)
    norm_ben = normalizar_texto(beneficiario)

    # 1. Procura pasta da comunidade
    pasta_com_real = None
    for item in doc_ativ.iterdir():
        if item.is_dir() and normalizar_texto(item.name) == norm_com:
            pasta_com_real = item
            break

    if not pasta_com_real:
        pasta_com_real = doc_ativ / (comunidade.upper() if comunidade else "GERAL")
        pasta_com_real.mkdir(parents=True, exist_ok=True)

    # 2. Procura pasta do beneficiário
    pasta_benef_real = None
    if pasta_com_real.exists():
        for b_dir in pasta_com_real.iterdir():
            if b_dir.is_dir() and normalizar_texto(b_dir.name) == norm_ben:
                pasta_benef_real = b_dir
                break

    # Se não achou na comunidade indicada, procura em todas as comunidades daquele técnico
    if not pasta_benef_real:
        for b_dir in doc_ativ.glob("*/*"):
            if b_dir.is_dir() and normalizar_texto(b_dir.name) == norm_ben:
                pasta_benef_real = b_dir
                break

    # Se ainda não existe, cria a pasta com o nome do beneficiário
    if not pasta_benef_real:
        pasta_benef_real = pasta_com_real / beneficiario.upper()
        pasta_benef_real.mkdir(parents=True, exist_ok=True)

    return pasta_benef_real

def organizar_planos(pasta_origem: Path, lista_autorizados: Any) -> Dict[str, Any]:
    """
    Processa e move os PDFs da pasta de origem SOMENTE SE o beneficiário
    estiver presente na lista de autorizados.
    """
    if not pasta_origem.exists():
        print(f"❌ Pasta de origem não encontrada: {pasta_origem}")
        return {"movidos": [], "ignorados": [], "erros": []}

    permitidos = carregar_lista_permitidos(lista_autorizados)
    print(f"📋 Lista de autorização carregada com {len(permitidos)} chaves (nomes/CPFs).")

    pdf_files = sorted(list(pasta_origem.glob("*.pdf")))
    print(f"📁 Encontrados {len(pdf_files)} arquivos PDF na pasta de origem.\n")

    movidos = []
    ignorados = []
    erros = []

    for pdf_path in pdf_files:
        try:
            dados = extrair_dados_pdf(pdf_path)
            nome_norm = normalizar_texto(dados["nome"])
            cpf_norm = normalizar_cpf(dados["cpf"])

            # VERIFICAÇÃO NA WHITELIST
            autorizado = False
            if nome_norm in permitidos or cpf_norm in permitidos:
                autorizado = True
            else:
                for perm in permitidos:
                    if len(perm) > 8 and (perm in nome_norm or nome_norm in perm):
                        autorizado = True
                        break

            if not autorizado:
                ignorados.append({
                    "arquivo": pdf_path.name,
                    "beneficiario": dados["nome"],
                    "cpf": dados["cpf"],
                    "comunidade": dados["comunidade"],
                    "motivo": "Não está na lista de autorizados"
                })
                continue

            # Localiza a pasta do técnico
            pasta_tecnico = encontrar_pasta_tecnico(dados["tecnico"])
            if not pasta_tecnico:
                for sub in BASE_TECNICOS_DIR.iterdir():
                    if sub.is_dir() and sub in pdf_path.parents:
                        pasta_tecnico = sub
                        break

            if not pasta_tecnico:
                erros.append({
                    "arquivo": pdf_path.name,
                    "beneficiario": dados["nome"],
                    "erro": f"Pasta do técnico '{dados['tecnico']}' não identificada."
                })
                continue

            # Localiza pasta do beneficiário e da comunidade
            pasta_benef = encontrar_ou_criar_pasta_beneficiario(
                pasta_tecnico, dados["comunidade"], dados["nome"]
            )

            # Cria pasta da atividade: DD.MM.AAAA - PLANO PRODUTIVO
            data_fmt = dados["data_formatada"]
            nome_pasta_ativ = f"{data_fmt} - PLANO PRODUTIVO"
            pasta_atividade = pasta_benef / nome_pasta_ativ
            pasta_atividade.mkdir(parents=True, exist_ok=True)

            # Nome final do arquivo
            nome_benef_pasta = pasta_benef.name
            nome_arquivo_final = f"{nome_benef_pasta} - COLLETUM.pdf"
            destino_final = pasta_atividade / nome_arquivo_final

            # Substitui arquivo antigo se existir
            if destino_final.exists():
                destino_final.unlink()

            shutil.move(str(pdf_path), str(destino_final))

            movidos.append({
                "arquivo_origem": pdf_path.name,
                "beneficiario": dados["nome"],
                "cpf": dados["cpf"],
                "destino": str(destino_final.relative_to(BASE_TECNICOS_DIR))
            })
            print(f"✅ [MOVIDO] {dados['nome']} -> {destino_final.relative_to(BASE_TECNICOS_DIR)}")

        except Exception as e:
            erros.append({
                "arquivo": pdf_path.name,
                "erro": str(e)
            })
            print(f"❌ [ERRO] {pdf_path.name}: {e}")

    print(f"\n" + "="*70)
    print(f"📊 RESUMO DO PROCESSAMENTO:")
    print(f"   • ✅ Movidos e Organizados com Sucesso: {len(movidos)}")
    print(f"   • ⏭️ Ignorados (Fora da Lista Autorizada): {len(ignorados)}")
    print(f"   • ❌ Erros: {len(erros)}")
    print(f"="*70)

    return {
        "movidos": movidos,
        "ignorados": ignorados,
        "erros": erros
    }

if __name__ == "__main__":
    print("Script organizador de Planos Produtivos pronto para uso!")
